"""生成带颜色的简历版本"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# 三个配色方案
THEMES = {
    "深蓝商务": {
        "primary": RGBColor(0x1B, 0x3A, 0x5C),   # 深蓝
        "accent": RGBColor(0x2B, 0x67, 0xB8),     # 亮蓝
        "line": RGBColor(0x2B, 0x67, 0xB8),
        "gray": RGBColor(0x66, 0x66, 0x66),
    },
    "墨绿自然": {
        "primary": RGBColor(0x1A, 0x4A, 0x3A),   # 墨绿
        "accent": RGBColor(0x2E, 0x8B, 0x6E),     # 翠绿
        "line": RGBColor(0x2E, 0x8B, 0x6E),
        "gray": RGBColor(0x66, 0x66, 0x66),
    },
    "棕红暖调": {
        "primary": RGBColor(0x5C, 0x2E, 0x1A),   # 棕红
        "accent": RGBColor(0xA0, 0x50, 0x30),     # 砖红
        "line": RGBColor(0xA0, 0x50, 0x30),
        "gray": RGBColor(0x66, 0x66, 0x66),
    },
}

def generate(theme_name, colors):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    BLACK = RGBColor(0, 0, 0)
    GRAY = colors["gray"]

    def sf(run, size=10.5, bold=False, color=BLACK):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def no_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            '<w:tcBorders %s>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>' % nsdecls("w")
        )
        tcPr.append(tcBorders)

    def add_section(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        # 左侧色块 + 标题
        r = p.add_run("  ")
        r.font.size = Pt(14)
        r.font.color.rgb = colors["line"]
        r2 = p.add_run(f"  {title}")
        sf(r2, size=14, bold=True, color=colors["primary"])
        # 下划线
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "8",
            qn("w:space"): "4", qn("w:color"): f"{colors['line'][0]:02X}{colors['line'][1]:02X}{colors['line'][2]:02X}",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def bullet(text, indent=0.5, prefix="- "):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(f"{prefix}{text}")
        sf(r, size=10.5)

    def project_block(title, date, tech, items, result):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        sf(r, size=11.5, bold=True, color=colors["primary"])
        r = p.add_run(f"    {date}")
        sf(r, size=10, color=GRAY)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"技术栈：{tech}")
        sf(r, size=9.5, color=GRAY)

        for item in items:
            bullet(item, indent=0.6)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.6)
        r = p.add_run(f"▸ {result}")
        sf(r, size=10.5, bold=True, color=colors["accent"])

    # ====== 头部（标签: 值 格式） ======
    info = [
        ("姓    名：", "李强"),
        ("教育背景：", "武汉工程科技学院 · 本科 · 2017-2021"),
        ("工作年限：", "4年测试经验（功能/接口/UI/AI）"),
        ("求职意向：", "AI软件测试工程师（上海/深圳/广州）"),
        ("期望薪资：", "15-20K · 1-2周到岗"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(2.8))
        r = p.add_run(label)
        sf(r, size=11, bold=True, color=colors["primary"])
        r = p.add_run(f"\t{value}")
        sf(r, size=11)

    # 个人优势单独一行（灰字）
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("独立搭建接口/UI 双端自动化回归体系  |  AI 测试工具开源可演示  |  量化评分驱动迭代至 82%")
    sf(r, size=9.5, color=GRAY)

    # ====== 技能亮点 ======
    add_section("核心优势")
    skills = [
        "AI辅助接口用例设计：基于大模型+量化评估闭环，单模块用例设计从1-2天压缩至3-5分钟，10模块平均评分82.0%",
        "手工用例智能生成：AI解析业务规则自动编排测试场景，8模块平均评分80.6%，步骤自动换行增强可读性",
        "接口自动化测试：独立搭建自动化回归体系，覆盖正向/反向/SQL注入等场景，回归时间缩短87.5%，累计发现7个后端缺陷",
        "测试设计方法论：等价类/边界值/判定表/场景法，覆盖正向、反向、异常、边界全场景",
    ]
    for s in skills:
        bullet(s)

    # ====== 项目经历 ======
    add_section("项目经历")
    project_block("AI大模型测试用例生成平台（个人开源项目）", "2025.04 - 至今",
        "deepseek-v4-flash, LangChain, Streamlit, DeepEval, Python", [
            "基于LangChain调用deepseek-v4-flash实现接口与手工用例智能生成，12种固定模板保底 + LLM动态生成深度业务场景",
            "手工用例支持AI解析业务规则自动编排测试步骤，动态列配置与数据驱动",
            "搭建DeepEval + 独立裁判模型双评估体系，覆盖数量达标率/字段完整性/断言规范性等8维度22模块",
            "基于Streamlit构建Web应用，支持项目/模块配置，一键导出Excel/JSON/Pytest脚本",
        ], "接口用例10模块平均评分82.0%，手工用例8模块平均80.6%；项目已开源可演示")
    project_block("接口自动化测试框架（教育平台）", "2025.01 - 至今",
        "Pytest, Requests, Jenkins, Allure, Python", [
            "独立搭建Pytest+Requests数据驱动接口自动化框架，初期覆盖登录/课程20+条核心用例，随版本迭代扩展至60+条，覆盖正向/反向/SQL注入等场景",
            "实现Token自动管理、动态测试数据隔离与自动清理，保障用例可随时重复执行",
            "对接Jenkins CI/CD，代码提交自动触发全量回归，推送Allure报告至项目组邮箱",
            "搭建分级日志系统，失败自动留存请求/响应快照，问题定位从30分钟缩短至5分钟",
        ], "回归时间从2小时降至15分钟（效率提升87.5%）；累计发现7个后端缺陷（含2个P1级）；支撑10+个版本顺利上线")
    project_block("B2C电商平台UI自动化测试框架", "2024.01 - 2024.12",
        "Playwright, Pytest, Jenkins, Allure, Python", [
            "基于Playwright+Pytest搭建POM分层UI自动化框架，JSON数据驱动覆盖登录/搜索/购物车/下单核心流程",
            "开发图片加载拦截使执行提速30%、登录态复用减少重复鉴权",
            "对接Jenkins CI/CD + Allure报告 + 邮件通知，提交代码自动执行并推送结果，实现无人值守回归",
            "随版本迭代持续维护自动化用例，分析自动化失败原因并跟踪缺陷至闭环",
        ], "页面维护成本降低60%，发现前端校验缺陷5处，全年支撑多次版本上线")
    project_block("电商平台功能测试", "2022.01 - 2023.12",
        "禅道, XMind, Postman", [
            "独立负责购物车/订单/搜索模块功能测试，运用等价类/边界值/判定表多方法设计用例80+条",
            "使用禅道全流程跟踪缺陷至闭环，发现8个缺陷（含1个P1级重复添加异常）",
        ], "发现8个缺陷（含1个P1级）")

    # ====== 工作经历 ======
    add_section("工作经历")
    works = [
        ("杭州砺信科技有限公司 | 测试工程师 | 2025.01 - 至今",
         "从0到1搭建接口自动化回归体系，回归效率提升87.5%，累计发现7个缺陷（含2个P1级），支撑10+个版本顺利上线；个人开源AI测试平台已可在线演示。"),
        ("杭州魂喵科技有限公司 | 测试工程师 | 2024.01 - 2024.12",
         "搭建电商UI自动化回归体系，页面维护成本降低60%，发现前端缺陷5个，落地Jenkins+Allure无人值守回归。"),
        ("杭州微风口网络科技有限公司 | 测试工程师 | 2022.01 - 2023.12",
         "负责电商核心模块功能测试，独立完成购物车用例设计与缺陷闭环，产出可复用测试资产。"),
    ]
    for t, d in works:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(t)
        sf(r, size=10.5, bold=True, color=colors["primary"])
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.6)
        r = p.add_run(d)
        sf(r, size=10.5)

    # ====== 自我评价 ======
    add_section("自我评价")
    evals = [
        "全栈测试能力：功能测试 + 接口自动化 + UI自动化 + AI测试，覆盖测试全流程",
        "结果导向：接口回归提效87.5%，UI维护成本降低60%，累计发现缺陷20+个",
        "AI工程化：大模型+测试设计结合，项目已开源可在线演示，具备量化评估体系",
        "快速适应：电商/教育/AI多领域项目经验，能快速融入新业务",
    ]
    for e in evals:
        bullet(e)

    # 保存
    out = f"李强_测试工程师简历_{theme_name}_final.docx"
    doc.save(out)
    print(f"Done: {out}")

for name, colors in THEMES.items():
    generate(name, colors)
