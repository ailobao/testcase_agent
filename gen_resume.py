"""生成 Word 简历"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

DARK = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x2B, 0x67, 0xB8)
GRAY = RGBColor(0x66, 0x66, 0x66)
DARK_TEXT = RGBColor(0x22, 0x22, 0x22)

def sf(run, name="微软雅黑", size=10.5, bold=False, color=DARK_TEXT):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.element.rPr.rFonts.set(qn("w:eastAsia"), name)

# ====== 头部 ======
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(15)
p.paragraph_format.space_after = Pt(2)
run = p.add_run("李  强")
sf(run, size=26, bold=True, color=DARK)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(0)
r = p2.add_run("27岁 | 本科 | 山西        测试工程师（上海/深圳/广州）        15-20K | 1-2周到岗")
sf(r, size=9.5, color=GRAY)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(6)
p3.paragraph_format.space_after = Pt(6)
r = p3.add_run("独立搭建接口/UI 双端自动化回归体系  |  AI 测试工具开源可演示  |  量化评分驱动迭代至 82%")
sf(r, size=9, color=RGBColor(0x55, 0x55, 0x55))

def add_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "8",
        qn("w:space"): "1", qn("w:color"): "2B67B8",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    sf(r, size=14, bold=True, color=DARK)
    add_line(doc)

def bullet(doc, text, bold_prefix="", indent=0.3, sz=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(indent)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        sf(r, size=sz, bold=True, color=DARK)
    r = p.add_run(text)
    sf(r, size=sz, color=DARK_TEXT)

# ====== 技能亮点 ======
add_section(doc, "技能亮点")
skills = [
    ("AI辅助接口用例设计", "基于大模型+量化评估闭环，单模块用例设计从1-2天压缩至3-5分钟，10模块平均评分82.0%"),
    ("手工用例智能生成", "AI解析业务规则自动编排测试场景，8模块平均评分80.6%，步骤自动换行增强可读性"),
    ("接口自动化测试", "独立搭建自动化回归体系，覆盖正向/反向/SQL注入等场景，回归时间缩短87.5%，累计发现7个后端缺陷"),
    ("测试设计方法论", "等价类/边界值/判定表/场景法，覆盖正向、反向、异常、边界全场景"),
]
for t, d in skills:
    bullet(doc, d, f"▸ {t}：")

# ====== 项目经历 ======
add_section(doc, "项目经历")

projects = [
    {
        "title": "AI大模型测试用例生成平台（个人开源项目）",
        "date": "2025.04 – 至今",
        "tech": "deepseek-v4-flash / LangChain / Streamlit / DeepEval / Python",
        "items": [
            "基于LangChain调用deepseek-v4-flash实现接口与手工用例智能生成，12种固定模板保底 + LLM动态生成深度业务场景",
            "手工用例支持AI解析业务规则自动编排测试步骤，动态列配置与数据驱动",
            "搭建DeepEval + 独立裁判模型双评估体系，覆盖数量达标率/字段完整性/断言规范性等8维度22模块",
            "基于Streamlit构建Web应用，支持项目/模块配置，一键导出Excel / JSON / Pytest脚本",
        ],
        "result": "接口用例10模块平均评分82.0%，手工用例8模块平均80.6%；项目已开源可演示",
    },
    {
        "title": "客达天下接口自动化测试框架（教育SaaS平台）",
        "date": "2025.01 – 至今",
        "tech": "Pytest / Requests / Jenkins / Allure / Python",
        "items": [
            "独立搭建Pytest+Requests数据驱动接口自动化框架，初期覆盖登录/课程20+条核心用例，随版本迭代扩展至60+条，覆盖正向/反向/SQL注入等场景",
            "实现Token自动管理、动态测试数据隔离与自动清理，保障用例可随时重复执行",
            "对接Jenkins CI/CD，代码提交自动触发全量回归，推送Allure报告至项目组邮箱",
            "搭建分级日志系统，失败自动留存请求/响应快照，问题定位从30分钟缩短至5分钟",
        ],
        "result": "回归时间从2小时降至15分钟（效率提升87.5%）；累计发现7个后端缺陷（含2个P1级）；支撑10+个版本顺利上线",
    },
    {
        "title": "B2C电商平台UI自动化测试框架",
        "date": "2024.01 – 2024.12",
        "tech": "Playwright / Pytest / Jenkins / Allure / Python",
        "items": [
            "基于Playwright+Pytest搭建POM分层UI自动化框架，JSON数据驱动覆盖登录/搜索/购物车/下单核心流程",
            "开发图片加载拦截使执行提速30%、登录态复用减少重复鉴权",
            "对接Jenkins CI/CD + Allure报告 + 邮件通知，提交代码自动执行并推送结果，实现无人值守回归",
            "随版本迭代持续维护自动化用例，分析自动化失败原因并跟踪缺陷至闭环",
        ],
        "result": "页面维护成本降低60%，发现前端校验缺陷5处，全年支撑多次版本上线",
    },
    {
        "title": "电商平台功能测试",
        "date": "2022.01 – 2023.12",
        "tech": "禅道 / XMind / Postman",
        "items": [
            "独立负责购物车/订单/搜索模块功能测试，运用等价类/边界值/判定表多方法设计用例80+条",
            "使用禅道全流程跟踪缺陷至闭环，发现8个缺陷（含1个P1级重复添加异常）",
            "输出《购物车模块测试Checklist》为后续版本提供复用依据",
        ],
        "result": "发现8个缺陷（含1个P1级），输出可复用测试Checklist",
    },
]

for proj in projects:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(proj["title"])
    sf(r, size=11, bold=True, color=DARK)
    r2 = p.add_run(f"    {proj['date']}")
    sf(r2, size=9, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f"技术栈：{proj['tech']}")
    sf(r, size=9, color=RGBColor(0x88, 0x88, 0x88))

    for item in proj["items"]:
        bullet(doc, item, indent=0.8, sz=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(f"成果：{proj['result']}")
    sf(r, size=10, color=ACCENT)

# ====== 工作经历 ======
add_section(doc, "工作经历")
works = [
    ("杭州砺信科技有限公司 | 测试工程师 | 2025.01 – 至今",
     "从0到1搭建接口自动化回归体系，回归效率提升87.5%，累计发现7个缺陷（含2个P1级），支撑10+个版本顺利上线；个人开源AI测试平台已可在线演示。"),
    ("杭州魂喵科技有限公司 | 测试工程师 | 2024.01 – 2024.12",
     "搭建电商UI自动化回归体系，页面维护成本降低60%，发现前端缺陷5个，落地Jenkins+Allure无人值守回归。"),
    ("杭州微风口网络科技有限公司 | 测试工程师 | 2022.01 – 2023.12",
     "负责电商核心模块功能测试，独立完成购物车用例设计与缺陷闭环，产出可复用测试资产。"),
]
for t, d in works:
    bullet(doc, "", f"▸ {t}", indent=0.3, sz=10.5)
    bullet(doc, d, indent=0.8, sz=10)

# ====== 自我评价 ======
add_section(doc, "自我评价")
evals = [
    ("全栈测试能力", "功能测试 + 接口自动化 + UI自动化 + AI测试，覆盖测试全流程"),
    ("结果导向", "接口回归提效87.5%，UI维护成本降低60%，累计发现缺陷20+个"),
    ("AI工程化", "大模型+测试设计结合，项目已开源可在线演示，具备量化评估体系"),
    ("快速适应", "电商/教育/AI多领域项目经验，能快速融入新业务"),
]
for t, d in evals:
    bullet(doc, d, f"▸ {t}：")

out = os.path.join(os.path.dirname(os.path.abspath(".")), "李强_测试工程师简历.docx")
doc.save(out)
print(f"Done: {out}")
